import docx
import logging
from typing import Dict, Union

class FileProcessor:
    """Class to handle processing of uploaded files"""
    
    def __init__(self):
        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
    
    def process_file(self, file_data: bytes, filename: str, category: str = None, country: str = None) -> Dict[str, Union[bool, str]]:
        """
        Process an uploaded file and extract its content
        
        Args:
            file_data (bytes): The file content in bytes
            filename (str): Name of the uploaded file
            category (str, optional): Brand category
            country (str, optional): Brand's primary country
            
        Returns:
            dict: Processed file data or error information
        """
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'txt':
                content = file_data.decode('utf-8')
            elif file_extension in ['doc', 'docx']:
                content = self._process_word_document(file_data)
            else:
                return {
                    "success": False,
                    "error": "Unsupported file type",
                    "message": "Please upload a .txt or .doc/.docx file"
                }
            
            # Create brand data object
            brand_name = filename.rsplit('.', 1)[0]  # Remove file extension
            brand_data = {
                "success": True,
                "brand_name": brand_name,
                "raw_text": content,
                "source_url": "File Upload",
                "category": category,
                "country": country
            }
            
            return brand_data
            
        except Exception as e:
            self.logger.error(f"Error processing file {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"File processing error: {str(e)}",
                "message": "An error occurred while processing the file."
            }
    
    def _process_word_document(self, file_data: bytes) -> str:
        """
        Extract text content from a Word document
        
        Args:
            file_data (bytes): The Word document content in bytes
            
        Returns:
            str: Extracted text content
        """
        try:
            from io import BytesIO
            doc = docx.Document(BytesIO(file_data))
            
            # Extract text from paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():  # Skip empty cells
                            full_text.append(cell.text)
            
            return '\n\n'.join(full_text)
            
        except Exception as e:
            self.logger.error(f"Error processing Word document: {str(e)}")
            raise 