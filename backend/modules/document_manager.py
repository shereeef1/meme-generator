import os
import json
from datetime import datetime
import docx
import logging
from typing import List, Dict, Optional

class DocumentManager:
    """Class to handle document storage, history, and management"""
    
    def __init__(self):
        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
        # Set up directories
        self.base_dir = os.path.dirname(os.path.dirname(__file__))
        self.docs_dir = os.path.join(self.base_dir, 'documents')
        self.history_file = os.path.join(self.docs_dir, 'history.json')
        
        # Create necessary directories
        os.makedirs(self.docs_dir, exist_ok=True)
        
        # Initialize history
        self._init_history()
    
    def _init_history(self):
        """Initialize or load document history"""
        if os.path.exists(self.history_file):
            try:
                with open(self.history_file, 'r') as f:
                    self.history = json.load(f)
            except:
                self.history = []
                self._save_history()
        else:
            self.history = []
            self._save_history()
    
    def _save_history(self):
        """Save document history to file"""
        with open(self.history_file, 'w') as f:
            json.dump(self.history, f, indent=2)
    
    def save_document_text(self, text_content: str, filename: str, category: str = None, country: str = None) -> str:
        """
        Save text content to a document file
        
        Args:
            text_content (str): Text content to save
            filename (str): Name for the document file
            category (str, optional): Brand category
            country (str, optional): Brand's primary country
            
        Returns:
            str: Path to the saved document
        """
        try:
            # Create a path for the document
            doc_path = os.path.join(self.docs_dir, filename)
            
            # Limit text size to prevent potential crashes
            MAX_TEXT_LENGTH = 50000  # Increased reasonable limit for Word document
            if len(text_content) > MAX_TEXT_LENGTH:
                self.logger.warning(f"Text content too large ({len(text_content)} chars), truncating to {MAX_TEXT_LENGTH} chars")
                text_content = text_content[:MAX_TEXT_LENGTH] + "\n\n[Content truncated due to size limitations...]"
            
            # Create a new Word document
            doc = docx.Document()
            
            # Add a title
            title = filename.rsplit('.', 1)[0].replace('_', ' ').title()
            doc.add_heading(title, 0)
            
            # Add metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if category:
                doc.add_paragraph(f"Category: {category}")
            if country:
                doc.add_paragraph(f"Country: {country}")
            
            # Add a separator
            doc.add_paragraph("=" * 50)
            
            # Add the content as paragraphs, with a size limit per paragraph
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Limit paragraph size to avoid issues with very long paragraphs
                    if len(para) > 5000:
                        para = para[:5000] + "... [paragraph truncated]"
                    
                    # Check if this is a section header
                    if para.startswith('===') and para.endswith('==='):
                        # Add as a heading
                        doc.add_heading(para.strip('=').strip(), 1)
                    else:
                        doc.add_paragraph(para.strip())
            
            # Save the document
            try:
                doc.save(doc_path)
                self.logger.info(f"Document saved: {doc_path}")
            except Exception as docx_error:
                self.logger.error(f"Error saving Word document: {str(docx_error)}")
                # If Word document save fails, immediately fall back to text file
                raise Exception(f"Word document save failed: {str(docx_error)}")
            
            # Add to history
            self.add_document(doc_path, f"Enhanced Research: {filename}", category, country)
            
            return doc_path
            
        except Exception as e:
            self.logger.error(f"Error saving document text: {str(e)}")
            self.logger.exception("Exception details:")
            
            # Create a fallback text file if Word document creation fails
            try:
                # Use a simple text file as fallback
                txt_filename = filename.rsplit('.', 1)[0] + '.txt'
                txt_path = os.path.join(self.docs_dir, txt_filename)
                
                self.logger.info(f"Attempting to save as text file instead: {txt_path}")
                
                with open(txt_path, 'w', encoding='utf-8') as f:
                    # Add a header to the text file
                    title = txt_filename.rsplit('.', 1)[0].replace('_', ' ').title()
                    f.write(f"{title}\n")
                    f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    if category:
                        f.write(f"Category: {category}\n")
                    if country:
                        f.write(f"Country: {country}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    # Limit text size for the text file too
                    if len(text_content) > 100000:
                        f.write(text_content[:100000] + "\n\n[Content truncated due to size limitations...]")
                    else:
                        f.write(text_content)
                
                # Add to history
                self.add_document(txt_path, f"Enhanced Research (Text): {txt_filename}", category, country)
                
                self.logger.info(f"Fallback text file saved: {txt_path}")
                return txt_path
            except Exception as fallback_error:
                self.logger.error(f"Even fallback save failed: {str(fallback_error)}")
                
                # Last resort - try to save a minimal text file with just the first part
                try:
                    minimal_filename = "minimal_" + txt_filename
                    minimal_path = os.path.join(self.docs_dir, minimal_filename)
                    
                    self.logger.info(f"Attempting minimal text save as last resort: {minimal_path}")
                    
                    with open(minimal_path, 'w', encoding='utf-8') as f:
                        f.write(f"EMERGENCY SAVE - DATA MAY BE INCOMPLETE\n")
                        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                        f.write("=" * 50 + "\n\n")
                        # Just save the first 10,000 characters
                        f.write(text_content[:10000] + "\n\n[Content severely truncated due to errors...]")
                    
                    self.add_document(minimal_path, f"Emergency Save: {minimal_filename}", category, country)
                    self.logger.info(f"Minimal emergency text file saved: {minimal_path}")
                    return minimal_path
                except:
                    self.logger.critical("ALL SAVE ATTEMPTS FAILED")
                    raise
    
    def add_document(self, doc_path: str, source_url: str, category: str = None, country: str = None) -> Dict:
        """
        Add a document to history
        
        Args:
            doc_path (str): Path to the document
            source_url (str): Original URL for scraped content
            category (str, optional): Brand category
            country (str, optional): Brand's primary country
            
        Returns:
            dict: Document entry
        """
        doc_name = os.path.basename(doc_path)
        entry = {
            "id": len(self.history) + 1,
            "filename": doc_name,
            "path": doc_path,
            "source_url": source_url,
            "category": category,
            "country": country,
            "created_at": datetime.now().isoformat(),
            "file_size": os.path.getsize(doc_path)
        }
        
        self.history.append(entry)
        self._save_history()
        return entry
    
    def get_document_history(self, limit: int = None, offset: int = 0) -> List[Dict]:
        """
        Get document history
        
        Args:
            limit (int, optional): Number of entries to return
            offset (int, optional): Number of entries to skip
            
        Returns:
            list: List of document entries
        """
        if limit:
            return self.history[offset:offset + limit]
        return self.history[offset:]
    
    def get_document(self, doc_id: int) -> Optional[Dict]:
        """
        Get document by ID
        
        Args:
            doc_id (int): Document ID
            
        Returns:
            dict: Document entry or None if not found
        """
        for entry in self.history:
            if entry["id"] == doc_id:
                return entry
        return None
    
    def update_document_content(self, doc_path: str, new_content: str) -> bool:
        """
        Update document content
        
        Args:
            doc_path (str): Path to the document
            new_content (str): New content to write
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Create new document
            doc = docx.Document()
            
            # Add content
            paragraphs = new_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save document
            doc.save(doc_path)
            return True
            
        except Exception as e:
            self.logger.error(f"Error updating document: {str(e)}")
            return False
    
    def delete_document(self, doc_id: int) -> bool:
        """
        Delete document by ID
        
        Args:
            doc_id (int): Document ID
            
        Returns:
            bool: True if successful, False otherwise
        """
        for i, entry in enumerate(self.history):
            if entry["id"] == doc_id:
                try:
                    # Delete file
                    if os.path.exists(entry["path"]):
                        os.remove(entry["path"])
                    
                    # Remove from history
                    self.history.pop(i)
                    self._save_history()
                    return True
                except Exception as e:
                    self.logger.error(f"Error deleting document: {str(e)}")
                    return False
        return False 