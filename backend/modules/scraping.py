import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import urlparse, urljoin
import logging
import traceback
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from collections import deque
import docx
from datetime import datetime
import os
import uuid
import urllib3

# Disable SSL warnings for requests
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

class BrandScraper:
    """Class to handle scraping of brand information from websites"""
    
    def __init__(self):
        # Set up logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
        # Set up Chrome options
        self.chrome_options = Options()
        self.chrome_options.add_argument('--headless')
        self.chrome_options.add_argument('--no-sandbox')
        self.chrome_options.add_argument('--disable-dev-shm-usage')
        self.chrome_options.add_argument('--disable-gpu')
        # Add timeout options
        self.chrome_options.add_argument('--page-load-timeout=30')
        self.chrome_options.add_argument('--script-timeout=30')
        # Ignore SSL errors and certificate issues
        self.chrome_options.add_argument('--ignore-certificate-errors')
        self.chrome_options.add_argument('--ignore-ssl-errors')
        self.chrome_options.add_argument('--allow-insecure-localhost')
        self.chrome_options.add_argument('--disable-web-security')
        
        # Initialize the webdriver
        self.service = Service(ChromeDriverManager().install())
        
        # Scraping settings
        self.max_pages = 20  # Reduced maximum pages to prevent timeouts
        self.delay = 1  # Delay between requests
        self.request_timeout = 30  # Timeout for requests in seconds
        
        # Create documents directory if it doesn't exist
        self.docs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'documents')
        os.makedirs(self.docs_dir, exist_ok=True)
    
    def scrape_brand_data(self, url, category=None, country=None):
        """
        Scrape brand data from a given URL and all its internal pages
        
        Args:
            url (str): The website URL to scrape
            category (str, optional): Brand category
            country (str, optional): Brand's primary country
            
        Returns:
            dict: Extracted brand data or error information
        """
        self.logger.info(f"Starting to scrape URL: {url}")
        driver = None
        try:
            # Validate URL
            if not self._is_valid_url(url):
                return {
                    "success": False,
                    "error": "Invalid URL",
                    "message": "Please provide a valid URL starting with http:// or https://"
                }
            
            # Get the domain name
            parsed_uri = urlparse(url)
            base_domain = parsed_uri.netloc
            if base_domain.startswith('www.'):
                base_domain = base_domain[4:]
            
            # Initialize the webdriver with explicit timeouts
            driver = webdriver.Chrome(service=self.service, options=self.chrome_options)
            driver.set_page_load_timeout(self.request_timeout)
            driver.set_script_timeout(self.request_timeout)
            
            # Initialize variables for crawling
            visited_urls = set()
            urls_to_visit = deque([url])
            all_content = []
            pages_scraped = 0
            scrape_success = False
            
            # Create Word document
            doc = docx.Document()
            doc.add_heading(f'Brand Information: {base_domain}', 0)
            
            # Add metadata
            doc.add_paragraph(f'Source URL: {url}')
            doc.add_paragraph(f'Scrape Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            if category:
                doc.add_paragraph(f'Category: {category}')
            if country:
                doc.add_paragraph(f'Country: {country}')
            
            doc.add_heading('Content', 1)
            
            # Try to access the homepage directly without waiting for page load
            try:
                self.logger.info(f"Attempting to access homepage: {url}")
                # First try with normal page load
                driver.get(url)
                
                try:
                    WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.TAG_NAME, "body"))
                    )
                    scrape_success = True
                except Exception as e:
                    self.logger.warning(f"Timeout waiting for page load: {str(e)}")
                    # Still continue as we might have partial content
                    
                # Parse homepage content even if waiting for full load fails
                soup = BeautifulSoup(driver.page_source, 'html.parser')
                page_text = self._extract_all_text(soup)
                
                if page_text:
                    self.logger.info("Successfully extracted text from homepage")
                    doc.add_paragraph(page_text)
                    all_content.append(f"\n=== Content from: {url} ===\n{page_text}")
                    scrape_success = True
                else:
                    self.logger.warning("No text content found on homepage")
                
            except Exception as e:
                self.logger.error(f"Failed to access homepage: {str(e)}")
                # We'll try to continue with other URLs
            
            # Crawl through all pages with more resilient error handling
            page_timeout_count = 0
            max_consecutive_failures = 3
            
            while urls_to_visit and pages_scraped < self.max_pages and page_timeout_count < max_consecutive_failures:
                current_url = urls_to_visit.popleft()
                if current_url in visited_urls:
                    continue
                    
                try:
                    self.logger.info(f"Scraping page: {current_url}")
                    driver.get(current_url)
                    time.sleep(self.delay)
                    
                    # Wait for page load with shorter timeout
                    try:
                        WebDriverWait(driver, 5).until(
                            EC.presence_of_element_located((By.TAG_NAME, "body"))
                        )
                        page_timeout_count = 0  # Reset consecutive failure count
                    except Exception as e:
                        self.logger.warning(f"Timeout waiting for page {current_url}: {str(e)}")
                        page_timeout_count += 1
                        # Continue anyway with whatever was loaded
                    
                    # Parse the page even if waiting for full load fails
                    soup = BeautifulSoup(driver.page_source, 'html.parser')
                    
                    # Add page URL to document
                    doc.add_heading(f'Content from: {current_url}', 2)
                    
                    # Extract and add text from this page
                    page_text = self._extract_all_text(soup)
                    if page_text:
                        doc.add_paragraph(page_text)
                        all_content.append(f"\n=== Content from: {current_url} ===\n{page_text}")
                        scrape_success = True  # If we got content from any page, count as success
                    
                    # Find all internal links (with error handling)
                    try:
                        links = soup.find_all('a', href=True)
                        for link in links:
                            href = link['href']
                            try:
                                full_url = urljoin(current_url, href)
                                # Only follow internal links
                                if self._is_internal_link(full_url, base_domain) and full_url not in visited_urls:
                                    urls_to_visit.append(full_url)
                            except Exception as link_error:
                                self.logger.warning(f"Error processing link {href}: {str(link_error)}")
                    except Exception as links_error:
                        self.logger.warning(f"Error finding links on {current_url}: {str(links_error)}")
                    
                    visited_urls.add(current_url)
                    pages_scraped += 1
                    self.logger.info(f"Pages scraped: {pages_scraped}")
                    
                except Exception as e:
                    self.logger.error(f"Error scraping page {current_url}: {str(e)}")
                    page_timeout_count += 1
                    # Add the URL to visited to avoid retrying
                    visited_urls.add(current_url)
                    continue
            
            # Check if we have any content at all
            if not scrape_success or not all_content:
                self.logger.warning(f"Failed to scrape any content from {url}")
                # Try a simplified approach as fallback
                try:
                    simplified_content = self._simplified_scrape(url)
                    if simplified_content:
                        self.logger.info("Successfully retrieved content using simplified scraping")
                        all_content = [simplified_content]
                        scrape_success = True
                        doc.add_paragraph(simplified_content)
                except Exception as simple_error:
                    self.logger.error(f"Simplified scraping also failed: {str(simple_error)}")
            
            # Generate unique filename
            filename = f"{base_domain}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{str(uuid.uuid4())[:8]}.docx"
            filepath = os.path.join(self.docs_dir, filename)
            
            # Save the document
            doc.save(filepath)
            
            # Combine all text
            combined_text = "\n".join(all_content)
            
            if not scrape_success or not combined_text.strip():
                # If we couldn't get any content, provide a basic fallback
                combined_text = f"Website: {url}\nBrand name: {base_domain}\nScraped on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                self.logger.warning("Using fallback text due to failed scraping")
            
            brand_data = {
                "success": True,  # Return success even if we only have fallback content
                "brand_name": base_domain,
                "tagline": self._extract_tagline(soup) if 'soup' in locals() else "",
                "description": self._extract_description(soup) if 'soup' in locals() else f"Data from {base_domain}",
                "products": self._extract_products(soup) if 'soup' in locals() else [],
                "category": category,
                "country": country,
                "source_url": url,
                "pages_scraped": pages_scraped,
                "raw_text": combined_text,
                "document_path": filepath,
                "document_name": filename,
                "scrape_quality": "full" if scrape_success else "fallback"
            }
            
            self.logger.info(f"Website crawling completed. Scraped {pages_scraped} pages.")
            return brand_data
            
        except Exception as e:
            self.logger.error(f"Scraping error: {str(e)}")
            self.logger.error(traceback.format_exc())
            
            # Create a basic fallback response
            try:
                parsed_uri = urlparse(url)
                base_domain = parsed_uri.netloc
                if base_domain.startswith('www.'):
                    base_domain = base_domain[4:]
                
                # Create a minimal document
                doc = docx.Document()
                doc.add_heading(f'Brand Information: {base_domain}', 0)
                doc.add_paragraph(f'Source URL: {url}')
                doc.add_paragraph(f'Scrape Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph("Note: Automated scraping failed. Basic information was created instead.")
                
                # Try to get some basic content using requests as a last resort
                fallback_content = self._simplified_scrape(url)
                if fallback_content:
                    doc.add_paragraph(fallback_content)
                
                # Save the document
                filename = f"{base_domain}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_fallback_{str(uuid.uuid4())[:8]}.docx"
                filepath = os.path.join(self.docs_dir, filename)
                doc.save(filepath)
                
                return {
                    "success": True,  # Return success with fallback data
                    "brand_name": base_domain,
                    "tagline": "",
                    "description": f"Basic information for {base_domain}",
                    "products": [],
                    "category": category,
                    "country": country,
                    "source_url": url,
                    "pages_scraped": 0,
                    "raw_text": fallback_content or f"Failed to scrape {url}. Please try uploading a document instead.",
                    "document_path": filepath,
                    "document_name": filename,
                    "scrape_quality": "minimal"
                }
            except Exception as fallback_error:
                self.logger.error(f"Even fallback creation failed: {str(fallback_error)}")
                return {
                    "success": False,
                    "error": f"Scraping error: {str(e)}",
                    "message": "An error occurred while processing the website content. Please try uploading a file instead."
                }
                
        finally:
            if driver:
                driver.quit()
                
    def _simplified_scrape(self, url):
        """Simple scraping using requests as a fallback"""
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            response = requests.get(url, headers=headers, timeout=10, verify=False)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(['script', 'style', 'noscript', 'iframe', 'head']):
                element.decompose()
                
            # Get main content
            content = ""
            
            # Try to get meta description
            meta_desc = soup.find('meta', {'name': 'description'})
            if meta_desc and meta_desc.get('content'):
                content += f"Description: {meta_desc.get('content')}\n\n"
                
            # Try to get main content areas
            for tag in ['main', 'article', 'section', 'div.content', 'div.main']:
                main_content = soup.select(tag)
                if main_content:
                    for element in main_content:
                        paragraphs = element.find_all('p')
                        for p in paragraphs:
                            if p.text.strip() and len(p.text.strip()) > 20:
                                content += p.text.strip() + "\n\n"
                                
            # If we couldn't find structured content, just get all paragraphs
            if not content:
                paragraphs = soup.find_all('p')
                for p in paragraphs:
                    if p.text.strip() and len(p.text.strip()) > 20:
                        content += p.text.strip() + "\n\n"
                        
            # If still no content, get all text as a last resort
            if not content:
                content = soup.get_text(separator="\n", strip=True)
                
            return content
        except Exception as e:
            self.logger.error(f"Simplified scraping failed: {str(e)}")
            return "Failed to retrieve content from website."
    
    def _is_internal_link(self, url, base_domain):
        """Check if a URL is an internal link"""
        try:
            parsed = urlparse(url)
            if not parsed.netloc:
                return True
            return base_domain in parsed.netloc
        except:
            return False
    
    def _is_valid_url(self, url):
        """Check if a URL is valid"""
        try:
            if not isinstance(url, str):
                self.logger.warning(f"Invalid URL type: {type(url)}")
                return False
                
            if not url.startswith('http://') and not url.startswith('https://'):
                self.logger.info(f"URL missing protocol, adding https://: {url}")
                url = 'https://' + url
                
            result = urlparse(url)
            valid = all([result.scheme, result.netloc]) and result.scheme in ['http', 'https']
            
            if not valid:
                self.logger.warning(f"Invalid URL structure: {url}")
                
            return valid
        except Exception as e:
            self.logger.error(f"Error validating URL {url}: {str(e)}")
            return False
    
    def _extract_domain(self, url):
        """Extract domain name from URL"""
        try:
            parsed_uri = urlparse(url)
            domain = parsed_uri.netloc
            if domain.startswith('www.'):
                domain = domain[4:]
            return domain.split('.')[0].capitalize()
        except:
            return ""
    
    def _extract_brand_name(self, soup, fallback_domain):
        """Extract the brand name from the website"""
        # Try to get from meta tags first (most reliable)
        meta_name = soup.find('meta', {'property': ['og:site_name', 'og:title']}) or \
                   soup.find('meta', {'name': 'application-name'})
        if meta_name and meta_name.get('content'):
            name = meta_name.get('content').strip()
            # Clean up common suffixes
            name = re.sub(r'[\|\-–—] .*$', '', name).strip()
            if name:
                return name

        # Try to get from structured data
        script_tags = soup.find_all('script', {'type': 'application/ld+json'})
        for script in script_tags:
            try:
                import json
                data = json.loads(script.string)
                if isinstance(data, dict):
                    # Look for organization or website name
                    if 'name' in data:
                        return data['name'].strip()
                    elif '@graph' in data:
                        for item in data['@graph']:
                            if isinstance(item, dict) and 'name' in item and \
                               ('@type' in item and item['@type'] in ['Organization', 'WebSite']):
                                return item['name'].strip()
            except:
                continue

        # Try to get from title tag
        if soup.title:
            title = soup.title.text.strip()
            # Remove common title endings
            title = re.sub(r'[\|\-–—] .*$', '', title).strip()
            # Remove common words
            title = re.sub(r'\b(home|welcome to|official site|official website)\b', '', title, flags=re.I).strip()
            if title:
                return title
        
        # Try to get from header or logo
        for selector in ['#logo', '.logo', '.brand', '.brand-name', '.site-title']:
            element = soup.select_one(selector)
            if element:
                # Try to get text or image alt
                if element.name == 'img' and element.get('alt'):
                    alt_text = element.get('alt').strip()
                    # Clean up alt text
                    alt_text = re.sub(r'\s*logo\s*', '', alt_text, flags=re.I).strip()
                    if alt_text:
                        return alt_text
                else:
                    text = element.get_text().strip()
                    if text:
                        return text

        # Try to find h1 that looks like a brand name (short and prominent)
        h1s = soup.find_all('h1')
        for h1 in h1s:
            text = h1.get_text().strip()
            # Brand names are typically short
            if text and len(text.split()) <= 4 and not re.search(r'welcome|about|contact', text, re.I):
                return text

        # Process the domain name as fallback
        domain_parts = fallback_domain.split('.')
        # Join words if the domain has multiple parts
        if len(domain_parts) > 1:
            return ' '.join(word.capitalize() for word in domain_parts)
        return fallback_domain.replace('-', ' ').replace('_', ' ').title()
    
    def _extract_tagline(self, soup):
        """Extract tagline or slogan from the website"""
        # Check header elements for short text that might be a tagline
        for tag in ['h1', 'h2', 'h3']:
            elements = soup.find_all(tag)
            for element in elements:
                text = element.text.strip()
                # Taglines are typically short
                if 3 < len(text.split()) < 12 and not re.search(r'navigation|menu|contact|about', text, re.I):
                    return text
        
        # Check meta description as fallback
        meta_desc = soup.find('meta', {'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            desc = meta_desc.get('content').strip()
            # Try to extract a short phrase from the beginning
            sentences = re.split(r'[.!?]', desc)
            if sentences and len(sentences[0].split()) < 15:
                return sentences[0].strip()
        
        return ""
    
    def _extract_description(self, soup):
        """Extract company description from the website"""
        # First try meta description
        meta_desc = soup.find('meta', {'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            return meta_desc.get('content').strip()
        
        # Try to find an about section
        about_section = soup.find('section', {'id': re.compile(r'about', re.I)}) or \
                        soup.find('div', {'id': re.compile(r'about', re.I)}) or \
                        soup.find('section', {'class': re.compile(r'about', re.I)}) or \
                        soup.find('div', {'class': re.compile(r'about', re.I)})
        
        if about_section:
            paragraphs = about_section.find_all('p')
            if paragraphs:
                # Use the longest paragraph as it's likely to be the main description
                longest = max(paragraphs, key=lambda p: len(p.text.strip()))
                return longest.text.strip()
        
        # Try to find any paragraph on the homepage
        paragraphs = soup.find_all('p')
        if paragraphs:
            # Filter out very short paragraphs and navigation/footer text
            valid_paragraphs = [p.text.strip() for p in paragraphs 
                                if len(p.text.strip()) > 50 
                                and not re.search(r'cookie|privacy|terms|copyright', p.text, re.I)]
            if valid_paragraphs:
                return max(valid_paragraphs, key=len)
        
        return ""
    
    def _extract_products(self, soup):
        """Extract product names from the website"""
        products = []
        
        # Look for product listings
        product_elements = soup.find_all(['div', 'article', 'li'], {'class': re.compile(r'product|item', re.I)})
        
        for element in product_elements[:10]:  # Limit to top 10 products
            product_name = element.find(['h2', 'h3', 'h4', 'a'])
            if product_name:
                name = product_name.text.strip()
                if name and len(name) < 100:  # Reasonable product name length
                    products.append(name)
        
        # If no products found, try to find them in the navigation
        if not products:
            nav = soup.find(['nav', 'ul'], {'class': re.compile(r'nav|menu', re.I)})
            if nav:
                links = nav.find_all('a')
                for link in links:
                    text = link.text.strip()
                    # Skip common navigation items
                    if text and not re.search(r'home|about|contact|cart|login|sign in', text, re.I):
                        products.append(text)
        
        return products[:10]  # Return up to 10 products
    
    def _infer_category(self, soup):
        """Try to infer the brand category from the website content"""
        text = soup.get_text().lower()
        
        # Define category keywords
        categories = {
            'fashion': ['clothing', 'apparel', 'wear', 'fashion', 'dress', 'shoe', 'accessory', 'style'],
            'beauty': ['beauty', 'cosmetic', 'makeup', 'skincare', 'fragrance', 'perfume'],
            'tech': ['technology', 'tech', 'electronics', 'digital', 'device', 'gadget', 'software'],
            'food': ['food', 'beverage', 'drink', 'snack', 'meal', 'recipe', 'cuisine'],
            'health': ['health', 'wellness', 'fitness', 'supplement', 'vitamin', 'nutrition'],
            'home': ['home', 'furniture', 'decor', 'kitchen', 'bedding', 'interior'],
        }
        
        # Count occurrences of keywords for each category
        scores = {}
        for category, keywords in categories.items():
            score = sum(1 for keyword in keywords if keyword in text)
            if score > 0:
                scores[category] = score
        
        # Return the category with the highest score, or None if no matches
        if scores:
            return max(scores.items(), key=lambda x: x[1])[0]
        
        return None
    
    def _extract_all_text(self, soup):
        """Extract all visible text from the webpage"""
        # Remove unwanted elements
        for element in soup(['script', 'style', 'noscript', 'iframe', 'head']):
            element.decompose()
        
        # Get text and clean it
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up the text
        lines = []
        for line in text.splitlines():
            line = line.strip()
            # Remove very short lines and common unwanted content
            if (len(line) > 5 and 
                not re.search(r'cookie|privacy|terms|copyright|all rights reserved', line, re.I) and
                not re.match(r'^[0-9\W]+$', line)):  # Skip lines with only numbers and special characters
                lines.append(line)
        
        cleaned_text = '\n'.join(lines)
        return cleaned_text 